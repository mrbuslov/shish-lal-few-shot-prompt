from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from bs4 import BeautifulSoup
from docx.shared import RGBColor
from typing import Union


class LocalDocxFormatter:
    """Class for formatting local docx files"""

    def __init__(self) -> None:
        pass

    def apply_html_formatting(self, paragraph: Paragraph, html_text: str) -> None:
        """Applies HTML formatting to the text"""
        soup = BeautifulSoup(html_text, "html.parser")

        # Get font from existing runs in paragraph
        base_font_name = None
        base_font_size = None
        for run in paragraph.runs:
            if run.font.name:
                base_font_name = run.font.name
            if run.font.size:
                base_font_size = run.font.size
            if base_font_name and base_font_size:
                break

        for elem in soup.descendants:
            if isinstance(elem, str) and elem.strip():
                new_run = paragraph.add_run(elem)

                # Apply base font from document
                if base_font_name:
                    new_run.font.name = base_font_name
                if base_font_size:
                    new_run.font.size = base_font_size

                # Looking for parent tags for applying styles
                parent = elem.parent
                while parent and parent.name:
                    if parent.name in ["b", "strong"]:
                        new_run.bold = True
                    if parent.name in ["i", "em"]:
                        new_run.italic = True
                    if parent.name == "u":
                        new_run.underline = True
                    if parent.name == "font":
                        color = parent.get("color")
                        if color and color.startswith("#") and len(color) == 7:
                            r, g, b = [int(color[i : i + 2], 16) for i in (1, 3, 5)]
                            new_run.font.color.rgb = RGBColor(r, g, b)
                    parent = parent.parent
            elif hasattr(elem, "name") and elem.name == "br":
                paragraph.add_run().add_break()

    def replace_text_with_formatting(
        self, paragraph: Paragraph, old: str, new: str, html: bool = False
    ) -> None:
        """Replaces text old with new in paragraph."""
        # Get all runs and their positions
        runs_to_process = list(paragraph.runs)

        for run_idx, run in enumerate(runs_to_process):
            if old not in run.text:
                continue

            parts = run.text.split(old)

            if len(parts) <= 1:
                continue

            # Save original formatting
            original_font_name = run.font.name
            original_font_size = run.font.size
            original_bold = run.bold
            original_italic = run.italic
            original_underline = run.underline
            original_color = None
            if run.font.color and run.font.color.rgb:
                original_color = run.font.color.rgb

            # Get position of current run in paragraph
            p_element = paragraph._element
            run_element = run._element
            run_position = list(p_element).index(run_element)

            # Set first part
            run.text = parts[0]

            # Insert replacements and remaining parts at correct position
            insert_position = run_position + 1
            for i in range(1, len(parts)):
                if html:
                    # For HTML, we need to insert at position
                    self.apply_html_formatting(paragraph, new)
                    # Move newly added runs to correct position
                    new_runs = paragraph.runs[len(runs_to_process) :]
                    for new_run in new_runs:
                        p_element.remove(new_run._element)
                        p_element.insert(insert_position, new_run._element)
                        insert_position += 1
                else:
                    # Create new run with original formatting
                    new_run = paragraph.add_run(new)
                    new_run.font.name = original_font_name
                    new_run.font.size = original_font_size
                    new_run.bold = original_bold
                    new_run.italic = original_italic
                    new_run.underline = original_underline
                    if original_color:
                        new_run.font.color.rgb = original_color

                    # Move to correct position
                    p_element.remove(new_run._element)
                    p_element.insert(insert_position, new_run._element)
                    insert_position += 1

                # Add remaining part with original formatting
                if parts[i]:
                    remaining_run = paragraph.add_run(parts[i])
                    remaining_run.font.name = original_font_name
                    remaining_run.font.size = original_font_size
                    remaining_run.bold = original_bold
                    remaining_run.italic = original_italic
                    remaining_run.underline = original_underline
                    if original_color:
                        remaining_run.font.color.rgb = original_color

                    # Move to correct position
                    p_element.remove(remaining_run._element)
                    p_element.insert(insert_position, remaining_run._element)
                    insert_position += 1

            # Update runs list for next iteration
            runs_to_process = list(paragraph.runs)

    def replace_all(
        self, doc: Union[Document, _Cell], old: str, new: str, html: bool = False
    ) -> None:
        """
        Goes through the entire document, including tables.

        Args:
            doc: python-docx Document object or _Cell object
            old: Text to replace
            new: New text (can be HTML if html=True)
            html: If True, new is interpreted as HTML
        """
        for paragraph in doc.paragraphs:
            self.replace_text_with_formatting(paragraph, old, new, html=html)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_all(cell, old, new, html=html)
